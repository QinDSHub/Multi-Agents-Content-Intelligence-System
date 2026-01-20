#!/usr/bin/env python
# coding: utf-8

from docx import Document

def normalize_list(x):
    if isinstance(x, list):
        return x
    if isinstance(x, str):
        return [i.strip() for i in x.split(",")]
    return 'Error'

# def export_to_word(columns, data, file_path, company_name):
#     doc = Document()
#     doc.add_heading("%s Content Strategy Plan"%company_name, 0)
#     for idx, item in enumerate(data):
#         item = normalize_list(item)
#         doc.add_heading(item[columns[0]], level=1)
#         doc.add_paragraph("Key Insights:"+'\n'+'\n'.join([x+';' for x in item[columns[1]]]))
#         doc.add_paragraph("Objective:"+'\n'+item[columns[2]])
#         doc.add_paragraph("Target Audience:"+'\n'+'\n'.join([x+';' for x in item[columns[3]].split(', ')]))
#         doc.add_paragraph("Recommended Distribution Strategy:"+'\n'+'\n'.join([x+';' for x in item[columns[4]].split(', ')]))
#         doc.add_paragraph("Priority Level (1=highest):"+'\n'+str(item[columns[5]]))
#         doc.add_paragraph("Explanation:"+'\n'+str(item[columns[6]]))
#     doc.save(file_path)
#     print(f"✅ Word file saved at: {file_path}")

def export_to_word(columns, data, file_path, company_name):
    doc = Document()
    doc.add_heading(f"{company_name} Content Strategy Plan", 0)
    for idx, item in enumerate(data):
        # 直接使用属性
        doc.add_heading(getattr(item, columns[0]), level=1)
        doc.add_paragraph("Key Insights:\n" + "\n".join(item.key_insights))
        doc.add_paragraph("Objective:\n" + item.target_objective)
        doc.add_paragraph("Target Audience:\n" + "\n".join(item.target_audience))
        doc.add_paragraph("Recommended Distribution Strategy:\n" + "\n".join(item.distribution_strategy))
        doc.add_paragraph("Priority Level (1=highest):\n" + str(item.priority))
        doc.add_paragraph("Explanation:\n" + item.explanation)
    doc.save(file_path)
    print(f"✅ Word file saved at: {file_path}")
